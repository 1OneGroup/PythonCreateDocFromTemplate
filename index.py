import docx
import comtypes.client
import pythoncom
import os
from flask import Flask, request ,send_file
import re 





app = Flask(__name__)

@app.route('/convert')
def convert_to_pdf():
    pythoncom.CoInitialize()

    # requestData = request.get_json()
    # print(requestData)
    file_path = os.path.join(os.getcwd(), 'OPR-PO Template.docx')
    file_pathpdf = os.path.join(os.getcwd(), 'file.pdf')

    doc = docx.Document(file_path)
    first_data={
        'Vendor':"chandan Nath",
        'Vendor ID':"dfdffdfd455",
        'Vendor Job Material':"shuttring"
    }
    for para in doc.paragraphs:
        matches = re.findall(r"{{(.*?)}}", para.text)
        for match in matches:
                    t=str(match)
                    try:
                        para.text = para.text.replace("{{"+t+"}}", first_data[t])
                    except:
                        para.text = para.text.replace("{{"+t+"}}", "error....")


    table = doc.tables[4]
    
    data = [{'Product Count': 1, 'Product': 'Product 1', 'Brand': 'Brand 1', 'Specification': 'Specification 1', 'Purchase Qty.': 6, 'Unit': 'Unit 1', 'Unit Price': 85.63691835145036, 'Net Value': 332.8074489935681, 'Tax Rate%': 1.6208329912419015, 'GST': 54.433213090697684, 'Total Amount': 1602.1599563791565}, {'Product Count': 2, 'Product': 'Product 2', 'Brand': 'Brand 2', 'Specification': 'Specification 2', 'Purchase Qty.': 8, 'Unit': 'Unit 2', 'Unit Price': 23.572334801123837, 'Net Value': 50.333982081612284, 'Tax Rate%': 3.6125884753619193, 'GST': 50.48044061112302, 'Total Amount': 2398.2061225350826}, {'Product Count': 3, 'Product': 'Product 3', 'Brand': 'Brand 3', 'Specification': 'Specification 3', 'Purchase Qty.': 9, 'Unit': 'Unit 3', 'Unit Price': 52.398389955204784, 'Net Value': 629.4658600886964, 'Tax Rate%': 5.079436106487902, 'GST': 67.99465220640441, 'Total Amount': 2873.8350254954953}, {'Product Count': 4, 'Product': 'Product 4', 'Brand': 'Brand 4', 'Specification': 'Specification 4', 'Purchase Qty.': 8, 'Unit': 'Unit 4', 'Unit Price': 62.459200467800684, 'Net Value': 856.9337858105595, 'Tax Rate%': 5.032637564576424, 'GST': 17.20517809184899, 'Total Amount': 7373.42960059784}, {'Product Count': 5, 'Product': 'Product 5', 'Brand': 'Brand 5', 'Specification': 'Specification 5', 'Purchase Qty.': 7, 'Unit': 'Unit 5', 'Unit Price': 97.0276142081059, 'Net Value': 254.6905122649565, 'Tax Rate%': 3.446089528189841, 'GST': 15.157648051409863, 'Total Amount': 214.9234904706282}]
    first_row=table.rows[0]

    for person in data:
        row = table.add_row()
        for index, cell in enumerate(first_row.cells):
                new_cell = row.cells[index]
                n=cell.text
                # tst=str(person[n])
                # new_cell.text =tst
                
                string = n
                matches = re.findall(r"{{(.*?)}}", string)
                tst=n
                for match in matches:
                    t=str(match)
                    # print(person["{{"+t+"}}"])
                    tst= tst.replace("{{"+t+"}}", str(person[t]))
                new_cell.text =tst




#    get the row you want to remove
    row_to_remove = table.rows[0]

    # remove the cells in the row
    for cell in row_to_remove.cells:
        cell._element.getparent().remove(cell._element)

    # remove the row element
    row_to_remove._element.getparent().remove(row_to_remove._element)




    for  index,table in enumerate(doc.tables):
        print(index)
        for row in table.rows:
            for cell in row.cells:
                # print(cell.text)
                
                matches = re.findall(r"{{(.*?)}}", cell.text)
                for match in matches:
                    t=str(match)
                    try:
                        cell.text = cell.text.replace("{{"+t+"}}", first_data[t])
                    except:
                        cell.text = cell.text.replace("{{"+t+"}}", "error....")
                        

    
    doc.save("updated_file.docx")

    vewfile_path = os.path.join(os.getcwd(), 'updated_file.docx')

    word = comtypes.client.CreateObject("Word.Application")
    word.Visible = False

    worddoc = word.Documents.Open(vewfile_path)
    worddoc.SaveAs(file_pathpdf, FileFormat=17)
    worddoc.Close()
    word.Quit()
    return send_file(file_pathpdf, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=False, host='0.0.0.0' )
